"""Conversation export -- MD, TXT, DOCX, PDF."""

from __future__ import annotations

import io
import json
import logging
import re
from datetime import datetime
from enum import Enum
from html.parser import HTMLParser
from typing import Any
from urllib.parse import quote

from fastapi import APIRouter, Depends, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from fim_agent.db import get_session
from fim_agent.web.auth import get_current_user
from fim_agent.web.exceptions import AppError
from fim_agent.web.models import Conversation, Message, User

router = APIRouter(prefix="/api/conversations", tags=["export"])
logger = logging.getLogger(__name__)


class ExportFormat(str, Enum):
    MD = "md"
    TXT = "txt"
    DOCX = "docx"
    PDF = "pdf"


class DetailLevel(str, Enum):
    FULL = "full"
    SUMMARY = "summary"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _sanitize_filename(title: str) -> str:
    """Replace characters unsafe for filenames with underscores."""
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", title)
    safe = safe.strip(". ")
    return safe or "conversation"


def _format_date(dt: datetime | None) -> str:
    """Return an ISO-8601 date string or empty string."""
    return dt.strftime("%Y-%m-%d %H:%M") if dt else ""


def _format_date_compact(dt: datetime | None) -> str:
    """Return a compact date for filenames like 20260309."""
    return dt.strftime("%Y%m%d") if dt else "export"


def _mode_label(mode: str) -> str:
    return "DAG (Plan & Execute)" if mode == "dag" else "Standard (ReAct)"


def _pair_messages(messages: list[Message]) -> list[dict[str, Any]]:
    """Group messages into turns: each turn has a user message and an
    optional assistant message.  Messages are expected to be sorted by
    ``created_at``."""
    turns: list[dict[str, Any]] = []
    current_turn: dict[str, Any] | None = None

    for msg in messages:
        if msg.role == "user":
            current_turn = {"user": msg, "assistant": None}
            turns.append(current_turn)
        elif msg.role == "assistant":
            if current_turn is not None:
                current_turn["assistant"] = msg
            else:
                # Orphan assistant message -- create a turn without user msg
                turns.append({"user": None, "assistant": msg})

    return turns


# Broad Unicode emoji ranges — covers SMP pictograph blocks, BMP symbol
# blocks commonly rendered with emoji presentation, and joiners/selectors.
_EMOJI_RE = re.compile(
    "["
    # ---- Supplemental Multilingual Plane (all pictograph blocks) ----
    "\U0001F000-\U0001FBFF"  # Mahjong → Playing Cards → Enclosed Alphanum Sup →
    #                          Enclosed Ideographic Sup → Misc Symbols & Pictographs →
    #                          Emoticons → Ornamental Dingbats → Transport/Map →
    #                          Supplemental Symbols → Geometric Shapes Ext →
    #                          Symbols Extended-A → Legacy Computing
    # ---- Basic Multilingual Plane – emoji-capable blocks ----
    "\U00002139"             # ℹ information source
    "\U00002194-\U00002199"  # ↔↕↖↗↘↙ arrows
    "\U000021A9-\U000021AA"  # ↩↪ curved arrows
    "\U0000231A-\U000023FF"  # Misc Technical (⌚⏰⏳⏺ etc.)
    "\U000024C2"             # Ⓜ circled M
    "\U000025A0-\U000025FF"  # Geometric Shapes (▶◻◼◽◾ etc.)
    "\U00002600-\U000027BF"  # Misc Symbols + Dingbats
    "\U00002934-\U00002935"  # ⤴⤵ curved arrows
    "\U00002B05-\U00002B07"  # ⬅⬆⬇ arrows
    "\U00002B1B-\U00002B1C"  # ⬛⬜ large squares
    "\U00002B50"             # ⭐ star
    "\U00002B55"             # ⭕ circle
    "\U00003030"             # 〰 wavy dash
    "\U0000303D"             # 〽 part alternation mark
    "\U00003297"             # ㊗ congratulation
    "\U00003299"             # ㊙ secret
    # ---- Joiners, modifiers, selectors ----
    "\U0000200B-\U0000200D"  # zero-width space / non-joiner / joiner
    "\U0000FE00-\U0000FE0F"  # variation selectors
    "\U000020E3"             # combining enclosing keycap
    "\U000E0020-\U000E007F"  # tags (flag subdivisions)
    "]+",
)


def _strip_emoji(text: str) -> str:
    """Remove emoji characters that DOCX/PDF fonts cannot render."""
    return _EMOJI_RE.sub("", text)


# ---------------------------------------------------------------------------
# SSE event parsing helpers
# ---------------------------------------------------------------------------


def _extract_sse_events(msg: Message) -> list[dict[str, Any]]:
    """Return the sse_events list from an assistant message's metadata."""
    meta = msg.metadata_ or {}
    return meta.get("sse_events", [])


def _extract_react_steps(events: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Extract completed tool-call steps from ReAct SSE events."""
    steps: list[dict[str, Any]] = []
    for ev in events:
        if ev.get("event") == "step":
            data = ev.get("data", {})
            if data.get("type") == "tool_call":
                steps.append(data)
    return steps


def _extract_done_event(events: list[dict[str, Any]]) -> dict[str, Any] | None:
    """Extract the done event from SSE events."""
    for ev in events:
        if ev.get("event") == "done":
            return ev.get("data", {})
    return None


def _extract_dag_plan(events: list[dict[str, Any]], target_round: int = 1) -> list[dict[str, Any]]:
    """Extract plan steps for a given DAG round."""
    for ev in events:
        if ev.get("event") == "phase":
            data = ev.get("data", {})
            if (
                data.get("name") == "planning"
                and data.get("status") == "done"
                and data.get("round", 1) == target_round
            ):
                return data.get("steps", [])
    return []


def _extract_dag_step_details(events: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    """Group step_progress events by step_id, collecting iterations and completion info."""
    steps: dict[str, dict[str, Any]] = {}
    for ev in events:
        if ev.get("event") == "step_progress":
            data = ev.get("data", {})
            sid = data.get("step_id", "")
            if sid not in steps:
                steps[sid] = {"iterations": [], "completed": None, "task": data.get("task", "")}

            event_type = data.get("event")
            if event_type == "iteration":
                steps[sid]["iterations"].append(data)
            elif event_type == "completed":
                steps[sid]["completed"] = data
            elif event_type == "started":
                steps[sid]["task"] = data.get("task", steps[sid]["task"])

    return steps


def _extract_dag_analysis(events: list[dict[str, Any]]) -> dict[str, Any] | None:
    """Extract the analysis phase result."""
    for ev in events:
        if ev.get("event") == "phase":
            data = ev.get("data", {})
            if data.get("name") == "analyzing" and data.get("status") == "done":
                return data
    return None


def _extract_dag_rounds(events: list[dict[str, Any]]) -> list[int]:
    """Determine which DAG rounds exist in the events."""
    rounds: set[int] = set()
    for ev in events:
        if ev.get("event") == "phase":
            data = ev.get("data", {})
            r = data.get("round")
            if r is not None:
                rounds.add(r)
    return sorted(rounds) if rounds else [1]


# ---------------------------------------------------------------------------
# Markdown renderer
# ---------------------------------------------------------------------------


def _render_md(conv: Conversation, messages: list[Message], detail: DetailLevel) -> str:
    """Render a conversation as a Markdown document."""
    lines: list[str] = []
    title = conv.title or "Untitled Conversation"

    # Header
    lines.append(f"# {title}")
    lines.append("")
    meta_parts = [
        f"**Mode:** {_mode_label(conv.mode)}",
    ]
    meta_parts.append(f"**Created:** {_format_date(conv.created_at)}")
    lines.append(" | ".join(meta_parts))

    if detail == DetailLevel.FULL and conv.total_tokens:
        lines.append(f"**Total Tokens:** {conv.total_tokens:,}")

    lines.append("")
    lines.append("---")
    lines.append("")

    # Turns
    turns = _pair_messages(messages)
    for idx, turn in enumerate(turns, 1):
        lines.append(f"## Turn {idx}")
        lines.append("")

        # User message
        user_msg: Message | None = turn.get("user")
        if user_msg:
            lines.append("**User:**")
            lines.append("")
            lines.append(user_msg.content or "")
            lines.append("")

        # Assistant message
        asst_msg: Message | None = turn.get("assistant")
        if asst_msg is None:
            lines.append("---")
            lines.append("")
            continue

        lines.append("**Assistant:**")
        lines.append("")

        if detail == DetailLevel.FULL:
            events = _extract_sse_events(asst_msg)
            done = _extract_done_event(events)

            if conv.mode == "dag":
                _render_md_dag_details(lines, events, done)
            else:
                _render_md_react_details(lines, events, done)

        # Final answer
        answer = asst_msg.content or ""
        if not answer and asst_msg.metadata_:
            answer = asst_msg.metadata_.get("answer", "")
        lines.append(answer)
        lines.append("")
        lines.append("---")
        lines.append("")

    return "\n".join(lines)


def _render_md_react_details(
    lines: list[str],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
) -> None:
    """Append ReAct execution details in Markdown."""
    steps = _extract_react_steps(events)
    if not steps:
        return

    iterations = done.get("iterations", len(steps)) if done else len(steps)
    elapsed = done.get("elapsed", 0) if done else 0

    lines.append(
        f"> **Execution Details** ({iterations} iteration{'s' if iterations != 1 else ''}, {elapsed:.1f}s)"
    )
    lines.append(">")
    lines.append("> | # | Tool | Duration |")
    lines.append("> |---|------|----------|")
    for i, step in enumerate(steps, 1):
        tool = step.get("tool_name", "unknown")
        dur = step.get("iter_elapsed", 0)
        lines.append(f"> | {i} | {tool} | {dur:.1f}s |")
    lines.append("")

    for i, step in enumerate(steps, 1):
        tool = step.get("tool_name", "unknown")
        dur = step.get("iter_elapsed", 0)
        lines.append(f"<details>")
        lines.append(f"<summary>Step {i}: {tool} ({dur:.1f}s)</summary>")
        lines.append("")

        reasoning = step.get("reasoning")
        if reasoning:
            lines.append(f"**Reasoning:** {reasoning}")
            lines.append("")

        args = step.get("tool_args")
        if args:
            lines.append("**Arguments:**")
            lines.append("```json")
            lines.append(json.dumps(args, indent=2, ensure_ascii=False))
            lines.append("```")
            lines.append("")

        observation = step.get("observation")
        if observation:
            lines.append("**Result:**")
            lines.append(str(observation))
            lines.append("")

        lines.append("</details>")
        lines.append("")


def _render_md_dag_details(
    lines: list[str],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
) -> None:
    """Append DAG execution details in Markdown."""
    rounds = _extract_dag_rounds(events)

    for rnd in rounds:
        plan_steps = _extract_dag_plan(events, rnd)
        if plan_steps:
            lines.append(f"> **Plan** (Round {rnd})")
            lines.append(">")
            lines.append("> | Step | Task | Dependencies |")
            lines.append("> |------|------|-------------|")
            for ps in plan_steps:
                sid = ps.get("id", "?")
                task = ps.get("task", "")
                deps = ", ".join(ps.get("deps", [])) or "\u2014"
                lines.append(f"> | {sid} | {task} | {deps} |")
            lines.append("")

    step_details = _extract_dag_step_details(events)
    for sid, info in step_details.items():
        task = info.get("task", "")
        completed = info.get("completed")
        status = completed.get("status", "completed") if completed else "unknown"
        duration = completed.get("duration", 0) if completed else 0

        lines.append("<details>")
        lines.append(f"<summary>{sid}: {task} ({status}, {duration:.1f}s)</summary>")
        lines.append("")

        iterations = info.get("iterations", [])
        if iterations:
            lines.append("**Iterations:**")
            for it in iterations:
                it_num = it.get("iteration", "?")
                tool = it.get("tool_name", "unknown")
                dur = it.get("iter_elapsed", 0)
                reasoning = it.get("reasoning", "")
                observation = it.get("observation", "")
                lines.append(
                    f"{it_num}. **{tool}** ({dur:.1f}s) \u2014 Reasoning: {reasoning}"
                )
                if observation:
                    lines.append(f"   Result: {observation}")
            lines.append("")

        result = completed.get("result", "") if completed else ""
        if result:
            lines.append(f"**Step Result:** {result}")
            lines.append("")

        lines.append("</details>")
        lines.append("")

    analysis = _extract_dag_analysis(events)
    if analysis:
        achieved = analysis.get("achieved", False)
        confidence = analysis.get("confidence", 0)
        reasoning = analysis.get("reasoning", "")
        achieved_label = "Goal Achieved" if achieved else "Goal Not Achieved"
        lines.append(
            f"> **Analysis:** {achieved_label} ({confidence * 100:.0f}% confidence)"
        )
        if reasoning:
            lines.append(f"> {reasoning}")
        lines.append("")


# ---------------------------------------------------------------------------
# TXT renderer
# ---------------------------------------------------------------------------


def _strip_md(text: str) -> str:
    """Naively strip Markdown formatting for plain-text output."""
    # Remove HTML tags
    text = re.sub(r"<[^>]+>", "", text)
    # Remove bold/italic markers
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    # Remove heading markers
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    # Remove code fences
    text = re.sub(r"```\w*\n?", "", text)
    # Remove horizontal rules
    text = re.sub(r"^---+\s*$", "=" * 60, text, flags=re.MULTILINE)
    return text


def _render_txt(conv: Conversation, messages: list[Message], detail: DetailLevel) -> str:
    """Render a conversation as plain text."""
    lines: list[str] = []
    title = conv.title or "Untitled Conversation"

    lines.append(title)
    lines.append("=" * len(title))
    lines.append("")
    lines.append(f"Mode: {_mode_label(conv.mode)}")
    lines.append(f"Created: {_format_date(conv.created_at)}")
    if detail == DetailLevel.FULL and conv.total_tokens:
        lines.append(f"Total Tokens: {conv.total_tokens:,}")
    lines.append("")
    lines.append("=" * 60)
    lines.append("")

    turns = _pair_messages(messages)
    for idx, turn in enumerate(turns, 1):
        lines.append(f"Turn {idx}")
        lines.append("-" * 40)
        lines.append("")

        user_msg: Message | None = turn.get("user")
        if user_msg:
            lines.append("[User]")
            lines.append("")
            lines.append(user_msg.content or "")
            lines.append("")

        asst_msg: Message | None = turn.get("assistant")
        if asst_msg is None:
            lines.append("=" * 60)
            lines.append("")
            continue

        lines.append("[Assistant]")
        lines.append("")

        if detail == DetailLevel.FULL:
            events = _extract_sse_events(asst_msg)
            done = _extract_done_event(events)

            if conv.mode == "dag":
                _render_txt_dag_details(lines, events, done)
            else:
                _render_txt_react_details(lines, events, done)

        answer = asst_msg.content or ""
        if not answer and asst_msg.metadata_:
            answer = asst_msg.metadata_.get("answer", "")
        lines.append(answer)
        lines.append("")
        lines.append("=" * 60)
        lines.append("")

    return "\n".join(lines)


def _render_txt_react_details(
    lines: list[str],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
) -> None:
    """Append ReAct details as plain text."""
    steps = _extract_react_steps(events)
    if not steps:
        return

    iterations = done.get("iterations", len(steps)) if done else len(steps)
    elapsed = done.get("elapsed", 0) if done else 0

    lines.append(
        f"  Execution Details: {iterations} iteration(s), {elapsed:.1f}s"
    )
    lines.append("")

    for i, step in enumerate(steps, 1):
        tool = step.get("tool_name", "unknown")
        dur = step.get("iter_elapsed", 0)
        lines.append(f"  Step {i}: {tool} ({dur:.1f}s)")

        reasoning = step.get("reasoning")
        if reasoning:
            lines.append(f"    Reasoning: {reasoning}")

        args = step.get("tool_args")
        if args:
            lines.append(f"    Arguments: {json.dumps(args, ensure_ascii=False)}")

        observation = step.get("observation")
        if observation:
            lines.append(f"    Result: {observation}")

        lines.append("")


def _render_txt_dag_details(
    lines: list[str],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
) -> None:
    """Append DAG details as plain text."""
    rounds = _extract_dag_rounds(events)

    for rnd in rounds:
        plan_steps = _extract_dag_plan(events, rnd)
        if plan_steps:
            lines.append(f"  Plan (Round {rnd}):")
            for ps in plan_steps:
                sid = ps.get("id", "?")
                task = ps.get("task", "")
                deps = ", ".join(ps.get("deps", [])) or "none"
                lines.append(f"    {sid}: {task} [deps: {deps}]")
            lines.append("")

    step_details = _extract_dag_step_details(events)
    for sid, info in step_details.items():
        task = info.get("task", "")
        completed = info.get("completed")
        status = completed.get("status", "completed") if completed else "unknown"
        duration = completed.get("duration", 0) if completed else 0

        lines.append(f"  {sid}: {task} ({status}, {duration:.1f}s)")

        for it in info.get("iterations", []):
            it_num = it.get("iteration", "?")
            tool = it.get("tool_name", "unknown")
            dur = it.get("iter_elapsed", 0)
            reasoning = it.get("reasoning", "")
            observation = it.get("observation", "")
            lines.append(f"    Iteration {it_num}: {tool} ({dur:.1f}s)")
            if reasoning:
                lines.append(f"      Reasoning: {reasoning}")
            if observation:
                lines.append(f"      Result: {observation}")

        result = completed.get("result", "") if completed else ""
        if result:
            lines.append(f"    Step Result: {result}")
        lines.append("")

    analysis = _extract_dag_analysis(events)
    if analysis:
        achieved = analysis.get("achieved", False)
        confidence = analysis.get("confidence", 0)
        reasoning = analysis.get("reasoning", "")
        label = "Goal Achieved" if achieved else "Goal Not Achieved"
        lines.append(f"  Analysis: {label} ({confidence * 100:.0f}% confidence)")
        if reasoning:
            lines.append(f"    {reasoning}")
        lines.append("")


# ---------------------------------------------------------------------------
# Markdown -> HTML helper (shared by DOCX and PDF renderers)
# ---------------------------------------------------------------------------


def _md_to_html(text: str) -> str:
    """Convert markdown to HTML using the markdown library."""
    import markdown

    return markdown.markdown(text, extensions=["fenced_code", "tables", "nl2br"])


# ---------------------------------------------------------------------------
# DOCX markdown renderer
# ---------------------------------------------------------------------------


class _DocxMarkdownRenderer(HTMLParser):
    """Parse HTML (converted from Markdown) and emit python-docx elements."""

    def __init__(self, doc: Any) -> None:
        super().__init__()
        self._doc = doc
        self._paragraph: Any | None = None
        self._bold = False
        self._italic = False
        self._code_inline = False
        self._in_pre = False
        self._pre_text = ""
        self._heading_level = 0
        self._heading_text = ""
        self._in_blockquote = False
        self._list_style: str | None = None
        self._list_stack: list[str] = []
        self._in_table = False
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] = []
        self._current_cell_text = ""
        self._in_td = False
        self._in_th = False
        self._href: str | None = None

    def _ensure_paragraph(self, style: str | None = None) -> Any:
        if self._paragraph is None:
            self._paragraph = self._doc.add_paragraph(style=style)
        return self._paragraph

    def _finish_paragraph(self) -> None:
        self._paragraph = None

    def _add_run(self, text: str) -> None:
        from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

        para = self._ensure_paragraph()
        run = para.add_run(text)
        if self._bold:
            run.bold = True
        if self._italic:
            run.italic = True
        if self._code_inline:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        if self._href:
            run.underline = True
            run.font.color.rgb = RGBColor(0x94, 0x6B, 0x2D)

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = dict(attrs)

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._heading_level = int(tag[1])
            self._heading_text = ""
        elif tag == "p":
            style = None
            if self._list_stack:
                style = self._list_stack[-1]
            self._paragraph = self._doc.add_paragraph(style=style)
            if self._in_blockquote:
                from docx.shared import Inches  # type: ignore[import-untyped]

                self._paragraph.paragraph_format.left_indent = Inches(0.5)
                self._italic = True
        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True
        elif tag == "code":
            if self._in_pre:
                pass  # handled by pre
            else:
                self._code_inline = True
        elif tag == "pre":
            self._in_pre = True
            self._pre_text = ""
        elif tag == "ul":
            self._list_stack.append("List Bullet")
        elif tag == "ol":
            self._list_stack.append("List Number")
        elif tag == "li":
            if self._list_stack:
                self._paragraph = self._doc.add_paragraph(style=self._list_stack[-1])
            else:
                self._paragraph = self._doc.add_paragraph(style="List Bullet")
        elif tag == "table":
            self._in_table = True
            self._table_rows = []
        elif tag == "tr":
            self._current_row = []
        elif tag in ("td", "th"):
            self._in_td = tag == "td"
            self._in_th = tag == "th"
            self._current_cell_text = ""
        elif tag == "blockquote":
            self._in_blockquote = True
        elif tag == "a":
            self._href = attrs_dict.get("href", "")
        elif tag == "hr":
            self._finish_paragraph()
            # Native DOCX horizontal rule via paragraph bottom border
            p = self._doc.add_paragraph()
            from docx.oxml.ns import qn  # type: ignore[import-untyped]
            from docx.oxml import OxmlElement  # type: ignore[import-untyped]

            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "C8B898")
            pBdr.append(bottom)
            pPr.append(pBdr)
            self._finish_paragraph()
        elif tag == "br":
            if self._paragraph:
                self._paragraph.add_run("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._doc.add_heading(self._heading_text, level=self._heading_level)
            self._heading_level = 0
            self._heading_text = ""
        elif tag == "p":
            if self._in_blockquote:
                self._italic = False
            self._finish_paragraph()
        elif tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag == "code":
            if not self._in_pre:
                self._code_inline = False
        elif tag == "pre":
            self._in_pre = False
            _add_monospace_paragraph(self._doc, self._pre_text)
            self._pre_text = ""
        elif tag == "ul":
            if self._list_stack:
                self._list_stack.pop()
        elif tag == "ol":
            if self._list_stack:
                self._list_stack.pop()
        elif tag == "li":
            self._finish_paragraph()
        elif tag in ("td", "th"):
            self._current_row.append(self._current_cell_text.strip())
            self._in_td = False
            self._in_th = False
            self._current_cell_text = ""
        elif tag == "tr":
            self._table_rows.append(self._current_row)
            self._current_row = []
        elif tag == "table":
            self._in_table = False
            if self._table_rows:
                self._emit_table()
        elif tag == "blockquote":
            self._in_blockquote = False
        elif tag == "a":
            self._href = None

    def handle_charref(self, name: str) -> None:
        """Convert HTML numeric character references (&#123; / &#xAB;) to
        characters and route through ``handle_data`` so emoji stripping
        applies uniformly."""
        try:
            codepoint = int(name[1:], 16) if name.startswith(("x", "X")) else int(name)
            self.handle_data(chr(codepoint))
        except (ValueError, OverflowError):
            pass

    def handle_data(self, data: str) -> None:
        data = _strip_emoji(data)
        if self._heading_level:
            self._heading_text += data
        elif self._in_pre:
            self._pre_text += data
        elif self._in_td or self._in_th:
            self._current_cell_text += data
        elif self._in_table:
            pass  # ignore whitespace between table elements
        else:
            if data.strip() or self._paragraph:
                self._add_run(data)

    def _emit_table(self) -> None:
        if not self._table_rows:
            return
        n_cols = max(len(r) for r in self._table_rows) if self._table_rows else 1
        table = self._doc.add_table(rows=len(self._table_rows), cols=n_cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(self._table_rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx < n_cols:
                    table.rows[r_idx].cells[c_idx].text = cell_text


def _md_to_docx(doc: Any, text: str) -> None:
    """Convert markdown text to DOCX elements via HTML intermediate."""
    html = _md_to_html(text)
    renderer = _DocxMarkdownRenderer(doc)
    renderer.feed(html)
    renderer.close()


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------


def _render_docx(conv: Conversation, messages: list[Message], detail: DetailLevel) -> bytes:
    """Render a conversation as a DOCX file and return the raw bytes."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise AppError(
            "docx_not_available",
            status_code=501,
            detail="DOCX export requires the python-docx package. "
            "Install with: uv pip install python-docx",
        )

    doc = Document()

    # Override default blue heading theme to match amber UI palette
    from docx.shared import RGBColor as _RGB
    _heading_color = _RGB(0x94, 0x6B, 0x2D)  # amber, matches UI primary
    for level in range(1, 7):
        style = doc.styles[f"Heading {level}"]
        style.font.color.rgb = _heading_color

    title = _strip_emoji(conv.title or "Untitled Conversation")

    # Title
    doc.add_heading(title, level=1)

    # Metadata
    meta_text = f"Mode: {_mode_label(conv.mode)}"
    meta_text += f"  |  Created: {_format_date(conv.created_at)}"
    if detail == DetailLevel.FULL and conv.total_tokens:
        meta_text += f"  |  Total Tokens: {conv.total_tokens:,}"
    doc.add_paragraph(meta_text)

    # Turns
    turns = _pair_messages(messages)
    for idx, turn in enumerate(turns, 1):
        doc.add_heading(f"Turn {idx}", level=2)

        user_msg: Message | None = turn.get("user")
        if user_msg:
            doc.add_heading("User", level=3)
            doc.add_paragraph(_strip_emoji(user_msg.content or ""))

        asst_msg: Message | None = turn.get("assistant")
        if asst_msg is None:
            continue

        doc.add_heading("Assistant", level=3)

        if detail == DetailLevel.FULL:
            events = _extract_sse_events(asst_msg)
            done = _extract_done_event(events)

            if conv.mode == "dag":
                _render_docx_dag_details(doc, events, done)
            else:
                _render_docx_react_details(doc, events, done)

        # Final answer
        answer = _strip_emoji(asst_msg.content or "")
        if not answer and asst_msg.metadata_:
            answer = _strip_emoji(asst_msg.metadata_.get("answer", ""))
        if answer:
            _md_to_docx(doc, answer)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_monospace_paragraph(doc: Any, text: str) -> None:
    """Add a paragraph with monospace font for code-like content."""
    try:
        from docx.shared import Pt
    except ImportError:
        doc.add_paragraph(text)
        return

    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def _render_docx_react_details(
    doc: Any,
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
) -> None:
    """Add ReAct execution details to a DOCX document."""
    steps = _extract_react_steps(events)
    if not steps:
        return

    iterations = done.get("iterations", len(steps)) if done else len(steps)
    elapsed = done.get("elapsed", 0) if done else 0

    doc.add_paragraph(
        f"Execution Details: {iterations} iteration(s), {elapsed:.1f}s"
    )

    for i, step in enumerate(steps, 1):
        tool = _strip_emoji(step.get("tool_name", "unknown"))
        dur = step.get("iter_elapsed", 0)

        doc.add_paragraph(
            f"Step {i}: {tool} ({dur:.1f}s)", style="List Bullet"
        )

        reasoning = step.get("reasoning")
        if reasoning:
            doc.add_paragraph(f"Reasoning: {_strip_emoji(reasoning)}", style="List Bullet 2")

        args = step.get("tool_args")
        if args:
            _add_monospace_paragraph(doc, _strip_emoji(json.dumps(args, indent=2, ensure_ascii=False)))

        observation = step.get("observation")
        if observation:
            doc.add_paragraph(f"Result: {_strip_emoji(str(observation)[:500])}", style="List Bullet 2")


def _render_docx_dag_details(
    doc: Any,
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
) -> None:
    """Add DAG execution details to a DOCX document."""
    rounds = _extract_dag_rounds(events)

    for rnd in rounds:
        plan_steps = _extract_dag_plan(events, rnd)
        if plan_steps:
            doc.add_paragraph(f"Plan (Round {rnd}):")
            for ps in plan_steps:
                sid = _strip_emoji(ps.get("id", "?"))
                task = _strip_emoji(ps.get("task", ""))
                deps = ", ".join(ps.get("deps", [])) or "none"
                doc.add_paragraph(
                    f"{sid}: {task} [deps: {deps}]", style="List Bullet"
                )

    step_details = _extract_dag_step_details(events)
    for sid, info in step_details.items():
        task = _strip_emoji(info.get("task", ""))
        completed = info.get("completed")
        status = _strip_emoji(completed.get("status", "completed") if completed else "unknown")
        duration = completed.get("duration", 0) if completed else 0

        doc.add_paragraph(f"{sid}: {task} ({status}, {duration:.1f}s)")

        for it in info.get("iterations", []):
            it_num = it.get("iteration", "?")
            tool = _strip_emoji(it.get("tool_name", "unknown"))
            dur = it.get("iter_elapsed", 0)
            reasoning = _strip_emoji(it.get("reasoning", ""))
            doc.add_paragraph(
                f"Iteration {it_num}: {tool} ({dur:.1f}s) - {reasoning}",
                style="List Bullet",
            )

            observation = it.get("observation", "")
            if observation:
                doc.add_paragraph(
                    f"Result: {_strip_emoji(str(observation)[:500])}", style="List Bullet 2"
                )

        result = completed.get("result", "") if completed else ""
        if result:
            doc.add_paragraph(f"Step Result: {_strip_emoji(result)}", style="List Bullet")

    analysis = _extract_dag_analysis(events)
    if analysis:
        achieved = analysis.get("achieved", False)
        confidence = analysis.get("confidence", 0)
        reasoning = _strip_emoji(analysis.get("reasoning", ""))
        label = "Goal Achieved" if achieved else "Goal Not Achieved"
        doc.add_paragraph(f"Analysis: {label} ({confidence * 100:.0f}% confidence)")
        if reasoning:
            doc.add_paragraph(reasoning)


# ---------------------------------------------------------------------------
# PDF renderer
# ---------------------------------------------------------------------------


def _register_cjk_font() -> str:
    """Register a CJK font for ReportLab and return the font name.

    Tries STSong-Light (Adobe CID font shipped with ReportLab);
    falls back to Helvetica if registration fails.
    """
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont

        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        logger.warning("CJK font STSong-Light unavailable, falling back to Helvetica")
        return "Helvetica"


def _build_pdf_styles(font_name: str) -> dict[str, Any]:
    """Build a dict of ReportLab paragraph styles for the PDF export."""
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch

    base = {
        "fontName": font_name,
        "wordWrap": "CJK",
        "alignment": TA_LEFT,
    }

    return {
        "title": ParagraphStyle(
            "pdf_title", fontSize=18, leading=24, spaceAfter=12,
            **base,
        ),
        "heading2": ParagraphStyle(
            "pdf_h2", fontSize=14, leading=18, spaceAfter=8, spaceBefore=12,
            **base,
        ),
        "heading3": ParagraphStyle(
            "pdf_h3", fontSize=12, leading=16, spaceAfter=6, spaceBefore=10,
            **base,
        ),
        "body": ParagraphStyle(
            "pdf_body", fontSize=10, leading=14, spaceAfter=6,
            **base,
        ),
        "code": ParagraphStyle(
            "pdf_code", fontName="Courier", fontSize=8, leading=10,
            spaceAfter=6, leftIndent=0.3 * inch,
            backColor="#F8F4ED", wordWrap="CJK",
        ),
        "meta": ParagraphStyle(
            "pdf_meta", fontSize=9, leading=12, spaceAfter=4,
            textColor="#6B5D4F", **base,
        ),
        "bullet": ParagraphStyle(
            "pdf_bullet", fontSize=10, leading=14, spaceAfter=4,
            leftIndent=0.4 * inch, bulletIndent=0.2 * inch,
            **base,
        ),
        "quote": ParagraphStyle(
            "pdf_quote", fontSize=10, leading=14, spaceAfter=6,
            leftIndent=0.4 * inch, textColor="#6B5D4F",
            fontName=font_name, wordWrap="CJK",
        ),
    }


class _PdfMarkdownRenderer(HTMLParser):
    """Parse HTML (converted from Markdown) and emit ReportLab flowables."""

    def __init__(self, styles: dict[str, Any], font_name: str) -> None:
        super().__init__()
        self._styles = styles
        self._font_name = font_name
        self.flowables: list[Any] = []

        self._text_buf = ""
        self._bold = False
        self._italic = False
        self._code_inline = False
        self._in_pre = False
        self._pre_text = ""
        self._heading_level = 0
        self._heading_text = ""
        self._in_blockquote = False
        self._in_li = False
        self._list_stack: list[str] = []  # "ul" or "ol"
        self._ol_counters: list[int] = []
        self._in_table = False
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] = []
        self._current_cell_text = ""
        self._in_td = False
        self._in_th = False
        self._href: str | None = None

    def _flush_text(self) -> None:
        """Flush accumulated text into a Paragraph flowable."""
        if not self._text_buf.strip():
            self._text_buf = ""
            return

        from reportlab.platypus import Paragraph

        style = self._styles["body"]
        if self._in_blockquote:
            style = self._styles["quote"]
        self.flowables.append(Paragraph(self._text_buf, style))
        self._text_buf = ""

    def _wrap_inline(self, text: str) -> str:
        """Wrap text with inline formatting tags for ReportLab Paragraph."""
        # Escape XML entities
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if self._bold:
            text = f"<b>{text}</b>"
        if self._italic:
            text = f"<i>{text}</i>"
        if self._code_inline:
            text = f'<font name="Courier" size="9">{text}</font>'
        if self._href:
            text = f'<u><font color="#946B2D">{text}</font></u>'
        return text

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        attrs_dict = dict(attrs)

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._flush_text()
            self._heading_level = int(tag[1])
            self._heading_text = ""
        elif tag == "p":
            pass  # text accumulates in _text_buf
        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True
        elif tag == "code":
            if not self._in_pre:
                self._code_inline = True
        elif tag == "pre":
            self._flush_text()
            self._in_pre = True
            self._pre_text = ""
        elif tag == "ul":
            self._flush_text()
            self._list_stack.append("ul")
        elif tag == "ol":
            self._flush_text()
            self._list_stack.append("ol")
            self._ol_counters.append(0)
        elif tag == "li":
            self._in_li = True
            self._text_buf = ""
        elif tag == "table":
            self._flush_text()
            self._in_table = True
            self._table_rows = []
        elif tag == "tr":
            self._current_row = []
        elif tag in ("td", "th"):
            self._in_td = tag == "td"
            self._in_th = tag == "th"
            self._current_cell_text = ""
        elif tag == "blockquote":
            self._flush_text()
            self._in_blockquote = True
        elif tag == "a":
            self._href = attrs_dict.get("href", "")
        elif tag == "hr":
            self._flush_text()
            from reportlab.platypus import HRFlowable

            self.flowables.append(HRFlowable(width="100%", thickness=0.5, color="#C8B898"))
        elif tag == "br":
            self._text_buf += "<br/>"

    def handle_endtag(self, tag: str) -> None:
        from reportlab.platypus import Paragraph, Preformatted, Spacer

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = self._heading_level
            style_key = "title" if level == 1 else ("heading2" if level <= 3 else "heading3")
            # Escape XML entities in heading text
            safe = self._heading_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            self.flowables.append(Paragraph(safe, self._styles[style_key]))
            self._heading_level = 0
            self._heading_text = ""
        elif tag == "p":
            self._flush_text()
        elif tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag == "code":
            if not self._in_pre:
                self._code_inline = False
        elif tag == "pre":
            self._in_pre = False
            safe = self._pre_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            self.flowables.append(Preformatted(safe, self._styles["code"]))
            self._pre_text = ""
        elif tag == "ul":
            if self._list_stack:
                self._list_stack.pop()
        elif tag == "ol":
            if self._list_stack:
                self._list_stack.pop()
            if self._ol_counters:
                self._ol_counters.pop()
        elif tag == "li":
            self._in_li = False
            text = self._text_buf.strip()
            self._text_buf = ""
            if text:
                prefix = "\u2022 "
                if self._list_stack and self._list_stack[-1] == "ol":
                    if self._ol_counters:
                        self._ol_counters[-1] += 1
                        prefix = f"{self._ol_counters[-1]}. "
                self.flowables.append(Paragraph(f"{prefix}{text}", self._styles["bullet"]))
        elif tag in ("td", "th"):
            self._current_row.append(self._current_cell_text.strip())
            self._in_td = False
            self._in_th = False
            self._current_cell_text = ""
        elif tag == "tr":
            self._table_rows.append(self._current_row)
            self._current_row = []
        elif tag == "table":
            self._in_table = False
            if self._table_rows:
                self._emit_table()
        elif tag == "blockquote":
            self._flush_text()
            self._in_blockquote = False
        elif tag == "a":
            self._href = None

    def handle_charref(self, name: str) -> None:
        """Convert HTML numeric character references to characters and route
        through ``handle_data`` so emoji stripping applies uniformly."""
        try:
            codepoint = int(name[1:], 16) if name.startswith(("x", "X")) else int(name)
            self.handle_data(chr(codepoint))
        except (ValueError, OverflowError):
            pass

    def handle_data(self, data: str) -> None:
        data = _strip_emoji(data)
        if self._heading_level:
            self._heading_text += data
        elif self._in_pre:
            self._pre_text += data
        elif self._in_td or self._in_th:
            self._current_cell_text += data
        elif self._in_table:
            pass
        elif self._in_li:
            self._text_buf += self._wrap_inline(data)
        else:
            self._text_buf += self._wrap_inline(data)

    def _emit_table(self) -> None:
        from reportlab.lib import colors
        from reportlab.platypus import Paragraph, Table, TableStyle

        if not self._table_rows:
            return

        n_cols = max(len(r) for r in self._table_rows)
        table_data: list[list[Any]] = []
        for row in self._table_rows:
            cells = []
            for i in range(n_cols):
                cell_text = row[i] if i < len(row) else ""
                safe = cell_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                cells.append(Paragraph(safe, self._styles["body"]))
            table_data.append(cells)

        if not table_data:
            return

        table = Table(table_data, repeatRows=1)
        table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.78, 0.72, 0.60)),
            ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.96, 0.93, 0.87)),
            ("FONTNAME", (0, 0), (-1, -1), self._font_name),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ]))
        self.flowables.append(table)


def _md_to_pdf_flowables(text: str, styles: dict[str, Any], font_name: str) -> list[Any]:
    """Convert markdown text to a list of ReportLab flowables via HTML intermediate."""
    html = _md_to_html(text)
    renderer = _PdfMarkdownRenderer(styles, font_name)
    renderer.feed(html)
    renderer.close()
    # Flush any remaining text
    renderer._flush_text()
    return renderer.flowables


def _render_pdf(conv: Conversation, messages: list[Message], detail: DetailLevel) -> bytes:
    """Render a conversation as a PDF file and return the raw bytes."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer
        from reportlab.platypus import HRFlowable
    except ImportError:
        raise AppError(
            "pdf_not_available",
            status_code=501,
            detail="PDF export requires the reportlab package. "
            "Install with: uv pip install reportlab",
        )

    font_name = _register_cjk_font()
    styles = _build_pdf_styles(font_name)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    flowables: list[Any] = []
    title = _strip_emoji(conv.title or "Untitled Conversation")

    # Title
    safe_title = title.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    flowables.append(Paragraph(safe_title, styles["title"]))

    # Metadata
    meta_text = f"Mode: {_mode_label(conv.mode)}"
    meta_text += f"  |  Created: {_format_date(conv.created_at)}"
    if detail == DetailLevel.FULL and conv.total_tokens:
        meta_text += f"  |  Total Tokens: {conv.total_tokens:,}"
    safe_meta = meta_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    flowables.append(Paragraph(safe_meta, styles["meta"]))
    flowables.append(Spacer(1, 12))
    flowables.append(HRFlowable(width="100%", thickness=1, color="#946B2D"))
    flowables.append(Spacer(1, 12))

    # Turns
    turns = _pair_messages(messages)
    for idx, turn in enumerate(turns, 1):
        flowables.append(Paragraph(f"Turn {idx}", styles["heading2"]))

        user_msg: Message | None = turn.get("user")
        if user_msg:
            flowables.append(Paragraph("User", styles["heading3"]))
            user_text = _strip_emoji(user_msg.content or "")
            safe_user = user_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flowables.append(Paragraph(safe_user, styles["body"]))
            flowables.append(Spacer(1, 6))

        asst_msg: Message | None = turn.get("assistant")
        if asst_msg is None:
            flowables.append(HRFlowable(width="100%", thickness=0.5, color="#C8B898"))
            flowables.append(Spacer(1, 8))
            continue

        flowables.append(Paragraph("Assistant", styles["heading3"]))

        if detail == DetailLevel.FULL:
            events = _extract_sse_events(asst_msg)
            done = _extract_done_event(events)

            if conv.mode == "dag":
                _render_pdf_dag_details(flowables, events, done, styles, font_name)
            else:
                _render_pdf_react_details(flowables, events, done, styles, font_name)

        # Final answer
        answer = _strip_emoji(asst_msg.content or "")
        if not answer and asst_msg.metadata_:
            answer = _strip_emoji(asst_msg.metadata_.get("answer", ""))
        if answer:
            answer_flowables = _md_to_pdf_flowables(answer, styles, font_name)
            flowables.extend(answer_flowables)

        flowables.append(Spacer(1, 6))
        flowables.append(HRFlowable(width="100%", thickness=0.5, color="#C8B898"))
        flowables.append(Spacer(1, 8))

    # Guard against empty document (ReportLab raises on no flowables)
    if not flowables:
        flowables.append(Paragraph("(empty conversation)", styles["body"]))

    doc.build(flowables)
    return buf.getvalue()


def _render_pdf_react_details(
    flowables: list[Any],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
    styles: dict[str, Any],
    font_name: str,
) -> None:
    """Add ReAct execution details as PDF flowables."""
    from reportlab.platypus import Paragraph, Preformatted, Spacer

    steps = _extract_react_steps(events)
    if not steps:
        return

    iterations = done.get("iterations", len(steps)) if done else len(steps)
    elapsed = done.get("elapsed", 0) if done else 0

    flowables.append(Paragraph(
        f"<b>Execution Details:</b> {iterations} iteration(s), {elapsed:.1f}s",
        styles["meta"],
    ))

    for i, step in enumerate(steps, 1):
        tool = _strip_emoji(step.get("tool_name", "unknown"))
        dur = step.get("iter_elapsed", 0)
        safe_tool = tool.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        flowables.append(Paragraph(
            f"\u2022 Step {i}: <b>{safe_tool}</b> ({dur:.1f}s)",
            styles["bullet"],
        ))

        reasoning = step.get("reasoning")
        if reasoning:
            safe_r = _strip_emoji(str(reasoning)).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flowables.append(Paragraph(f"Reasoning: {safe_r}", styles["quote"]))

        args = step.get("tool_args")
        if args:
            args_text = _strip_emoji(json.dumps(args, indent=2, ensure_ascii=False))
            safe_args = args_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flowables.append(Preformatted(safe_args, styles["code"]))

        observation = step.get("observation")
        if observation:
            obs_text = _strip_emoji(str(observation)[:500])
            safe_obs = obs_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flowables.append(Paragraph(f"Result: {safe_obs}", styles["quote"]))

    flowables.append(Spacer(1, 8))


def _render_pdf_dag_details(
    flowables: list[Any],
    events: list[dict[str, Any]],
    done: dict[str, Any] | None,
    styles: dict[str, Any],
    font_name: str,
) -> None:
    """Add DAG execution details as PDF flowables."""
    from reportlab.lib import colors
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

    rounds = _extract_dag_rounds(events)

    for rnd in rounds:
        plan_steps = _extract_dag_plan(events, rnd)
        if plan_steps:
            flowables.append(Paragraph(f"<b>Plan (Round {rnd}):</b>", styles["meta"]))

            # Build table: Step | Task | Dependencies
            header = [
                Paragraph("<b>Step</b>", styles["body"]),
                Paragraph("<b>Task</b>", styles["body"]),
                Paragraph("<b>Dependencies</b>", styles["body"]),
            ]
            table_data = [header]
            for ps in plan_steps:
                sid = _strip_emoji(ps.get("id", "?"))
                task = _strip_emoji(ps.get("task", ""))
                deps = ", ".join(ps.get("deps", [])) or "\u2014"
                safe_sid = str(sid).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                safe_task = task.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                safe_deps = deps.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                table_data.append([
                    Paragraph(safe_sid, styles["body"]),
                    Paragraph(safe_task, styles["body"]),
                    Paragraph(safe_deps, styles["body"]),
                ])

            table = Table(table_data, repeatRows=1)
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.78, 0.72, 0.60)),
                ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.96, 0.93, 0.87)),
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            flowables.append(table)
            flowables.append(Spacer(1, 8))

    step_details = _extract_dag_step_details(events)
    for sid, info in step_details.items():
        task = _strip_emoji(info.get("task", ""))
        completed = info.get("completed")
        status = _strip_emoji(completed.get("status", "completed") if completed else "unknown")
        duration = completed.get("duration", 0) if completed else 0
        safe_sid = _strip_emoji(str(sid)).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        safe_task = task.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        safe_status = status.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        flowables.append(Paragraph(
            f"<b>{safe_sid}:</b> {safe_task} ({safe_status}, {duration:.1f}s)",
            styles["body"],
        ))

        for it in info.get("iterations", []):
            it_num = it.get("iteration", "?")
            tool = _strip_emoji(it.get("tool_name", "unknown"))
            dur = it.get("iter_elapsed", 0)
            reasoning = _strip_emoji(it.get("reasoning", ""))
            observation = it.get("observation", "")
            safe_tool = str(tool).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            safe_reason = str(reasoning).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

            flowables.append(Paragraph(
                f"\u2022 Iteration {it_num}: <b>{safe_tool}</b> ({dur:.1f}s) \u2014 {safe_reason}",
                styles["bullet"],
            ))

            if observation:
                obs_text = _strip_emoji(str(observation)[:500])
                safe_obs = obs_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                flowables.append(Paragraph(f"Result: {safe_obs}", styles["quote"]))

        result = completed.get("result", "") if completed else ""
        if result:
            safe_result = _strip_emoji(str(result)).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flowables.append(Paragraph(f"Step Result: {safe_result}", styles["bullet"]))
        flowables.append(Spacer(1, 4))

    analysis = _extract_dag_analysis(events)
    if analysis:
        achieved = analysis.get("achieved", False)
        confidence = analysis.get("confidence", 0)
        reasoning = _strip_emoji(analysis.get("reasoning", ""))
        label = "Goal Achieved" if achieved else "Goal Not Achieved"
        flowables.append(Paragraph(
            f"<b>Analysis:</b> {label} ({confidence * 100:.0f}% confidence)",
            styles["meta"],
        ))
        if reasoning:
            safe_reason = str(reasoning).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flowables.append(Paragraph(safe_reason, styles["quote"]))
        flowables.append(Spacer(1, 8))


# ---------------------------------------------------------------------------
# Endpoint
# ---------------------------------------------------------------------------

_CONTENT_TYPES: dict[ExportFormat, str] = {
    ExportFormat.MD: "text/markdown; charset=utf-8",
    ExportFormat.TXT: "text/plain; charset=utf-8",
    ExportFormat.DOCX: (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document"
    ),
    ExportFormat.PDF: "application/pdf",
}


@router.get("/{conversation_id}/export")
async def export_conversation(
    conversation_id: str,
    format: ExportFormat = Query(..., description="Export format: md, txt, docx, or pdf"),
    detail: DetailLevel = Query(
        DetailLevel.FULL, description="Detail level: full or summary"
    ),
    current_user: User = Depends(get_current_user),  # noqa: B008
    db: AsyncSession = Depends(get_session),  # noqa: B008
) -> StreamingResponse:
    """Export a conversation as a downloadable file.

    Supports Markdown, plain text, DOCX, and PDF formats.  The ``detail``
    parameter controls whether tool execution details are included
    (``full``) or only the final answers (``summary``).
    """
    # Fetch conversation with messages, verify ownership
    result = await db.execute(
        select(Conversation)
        .where(
            Conversation.id == conversation_id,
            Conversation.user_id == current_user.id,
        )
        .options(selectinload(Conversation.messages))
    )
    conv = result.scalar_one_or_none()
    if conv is None:
        raise AppError("conversation_not_found", status_code=404)

    # Sort messages chronologically
    messages = sorted(conv.messages, key=lambda m: m.created_at)

    # Build filename
    safe_title = _sanitize_filename(conv.title or "conversation")
    date_str = _format_date_compact(conv.created_at)
    ext = format.value
    detail_suffix = "_full" if detail == DetailLevel.FULL else ""
    filename = f"{safe_title}_{date_str}{detail_suffix}.{ext}"

    # Render content
    if format == ExportFormat.PDF:
        content_bytes = _render_pdf(conv, messages, detail)
        stream = io.BytesIO(content_bytes)
    elif format == ExportFormat.DOCX:
        content_bytes = _render_docx(conv, messages, detail)
        stream = io.BytesIO(content_bytes)
    elif format == ExportFormat.TXT:
        text = _render_txt(conv, messages, detail)
        stream = io.BytesIO(text.encode("utf-8"))
    else:  # MD
        text = _render_md(conv, messages, detail)
        stream = io.BytesIO(text.encode("utf-8"))

    content_type = _CONTENT_TYPES[format]

    # RFC 5987: use ASCII fallback + UTF-8 encoded filename for non-ASCII titles
    ascii_filename = filename.encode("ascii", errors="ignore").decode("ascii") or f"export.{ext}"
    utf8_filename = quote(filename)
    disposition = f"attachment; filename=\"{ascii_filename}\"; filename*=UTF-8''{utf8_filename}"

    return StreamingResponse(
        stream,
        media_type=content_type,
        headers={
            "Content-Disposition": disposition,
        },
    )
