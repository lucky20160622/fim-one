"""Tests for conversation export renderers (MD, TXT, DOCX, PDF)."""

from __future__ import annotations

import io
import json
from datetime import datetime, timezone
from typing import Any
from unittest.mock import MagicMock

import pytest


# ---------------------------------------------------------------------------
# Helpers to build fake ORM objects without touching the DB
# ---------------------------------------------------------------------------

def _make_conv(
    *,
    title: str = "Test Conversation",
    mode: str = "react",
    model_name: str = "gpt-4o",
    total_tokens: int = 1500,
    created_at: datetime | None = None,
) -> MagicMock:
    conv = MagicMock()
    conv.id = "conv-001"
    conv.title = title
    conv.mode = mode
    conv.model_name = model_name
    conv.total_tokens = total_tokens
    conv.created_at = created_at or datetime(2026, 3, 9, 10, 30, 0)
    conv.user_id = "user-001"
    conv.messages = []
    return conv


def _make_msg(
    role: str,
    content: str,
    metadata_: dict[str, Any] | None = None,
    created_at: datetime | None = None,
) -> MagicMock:
    msg = MagicMock()
    msg.role = role
    msg.content = content
    msg.metadata_ = metadata_
    msg.created_at = created_at or datetime(2026, 3, 9, 10, 30, 0)
    return msg


def _react_events(steps: list[dict[str, Any]], answer: str = "Final.") -> list[dict[str, Any]]:
    """Build SSE events for a ReAct assistant message."""
    events: list[dict[str, Any]] = []
    for i, s in enumerate(steps, 1):
        events.append({
            "event": "step",
            "data": {
                "type": "tool_call",
                "iteration": i,
                "tool_name": s.get("tool", "web_search"),
                "tool_args": s.get("args", {"query": "test"}),
                "reasoning": s.get("reasoning", "need to search"),
                "observation": s.get("observation", "result data"),
                "error": None,
                "iter_elapsed": s.get("elapsed", 1.0),
            },
        })
    events.append({
        "event": "done",
        "data": {
            "answer": answer,
            "iterations": len(steps),
            "elapsed": sum(s.get("elapsed", 1.0) for s in steps),
            "usage": {"prompt_tokens": 100, "completion_tokens": 50, "total_tokens": 150},
        },
    })
    return events


def _dag_events(
    plan_steps: list[dict[str, Any]],
    answer: str = "DAG answer.",
) -> list[dict[str, Any]]:
    """Build SSE events for a DAG assistant message."""
    events: list[dict[str, Any]] = []
    # Planning phase
    events.append({"event": "phase", "data": {"name": "planning", "status": "start", "round": 1}})
    events.append({
        "event": "phase",
        "data": {
            "name": "planning",
            "status": "done",
            "round": 1,
            "steps": plan_steps,
        },
    })
    # Executing phase
    events.append({"event": "phase", "data": {"name": "executing", "status": "start", "round": 1}})
    for ps in plan_steps:
        sid = ps["id"]
        events.append({"event": "step_progress", "data": {"step_id": sid, "event": "started", "task": ps["task"]}})
        events.append({
            "event": "step_progress",
            "data": {
                "step_id": sid,
                "event": "iteration",
                "type": "tool_call",
                "iteration": 1,
                "tool_name": ps.get("tool_hint", "web_search"),
                "tool_args": {"query": "test"},
                "reasoning": "searching",
                "observation": "found data",
                "iter_elapsed": 1.2,
            },
        })
        events.append({
            "event": "step_progress",
            "data": {
                "step_id": sid,
                "event": "completed",
                "status": "completed",
                "result": f"Result for {sid}",
                "duration": 2.0,
            },
        })
    events.append({"event": "phase", "data": {"name": "executing", "status": "done", "round": 1}})
    # Analysis
    events.append({
        "event": "phase",
        "data": {
            "name": "analyzing",
            "status": "done",
            "round": 1,
            "achieved": True,
            "confidence": 0.9,
            "reasoning": "All tasks completed successfully.",
        },
    })
    # Done
    events.append({
        "event": "done",
        "data": {
            "answer": answer,
            "achieved": True,
            "confidence": 0.9,
            "elapsed": 5.0,
            "rounds": 1,
            "usage": {"prompt_tokens": 200, "completion_tokens": 100, "total_tokens": 300},
        },
    })
    return events


# ---------------------------------------------------------------------------
# Import the module under test
# ---------------------------------------------------------------------------

from fim_agent.web.api.export import (
    DetailLevel,
    ExportFormat,
    _extract_dag_analysis,
    _extract_dag_plan,
    _extract_dag_rounds,
    _extract_dag_step_details,
    _extract_done_event,
    _extract_react_steps,
    _extract_sse_events,
    _format_date,
    _format_date_compact,
    _md_to_docx,
    _mode_label,
    _pair_messages,
    _render_docx,
    _render_md,
    _render_pdf,
    _render_txt,
    _sanitize_filename,
    _strip_emoji,
)


# ===================================================================
# Unit tests: helper functions
# ===================================================================


class TestSanitizeFilename:
    def test_normal_title(self):
        assert _sanitize_filename("My Chat") == "My Chat"

    def test_special_chars(self):
        assert _sanitize_filename('File: "test" <1>') == "File_ _test_ _1_"

    def test_empty(self):
        assert _sanitize_filename("") == "conversation"

    def test_dots_and_spaces(self):
        assert _sanitize_filename("...") == "conversation"


class TestFormatDate:
    def test_none(self):
        assert _format_date(None) == ""

    def test_valid(self):
        dt = datetime(2026, 3, 9, 14, 30)
        assert _format_date(dt) == "2026-03-09 14:30"


class TestFormatDateCompact:
    def test_none(self):
        assert _format_date_compact(None) == "export"

    def test_valid(self):
        dt = datetime(2026, 1, 15, 8, 0)
        assert _format_date_compact(dt) == "20260115"


class TestModeLabel:
    def test_react(self):
        assert _mode_label("react") == "Standard"

    def test_dag(self):
        assert _mode_label("dag") == "Planner"

    def test_react_zh(self):
        assert _mode_label("react", "zh") == "标准"

    def test_dag_zh(self):
        assert _mode_label("dag", "zh") == "规划"


class TestPairMessages:
    def test_empty(self):
        assert _pair_messages([]) == []

    def test_single_user(self):
        msgs = [_make_msg("user", "Hello")]
        turns = _pair_messages(msgs)
        assert len(turns) == 1
        assert turns[0]["user"].content == "Hello"
        assert turns[0]["assistant"] is None

    def test_user_assistant_pair(self):
        msgs = [
            _make_msg("user", "Q1", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg("assistant", "A1", created_at=datetime(2026, 1, 1, 0, 1)),
        ]
        turns = _pair_messages(msgs)
        assert len(turns) == 1
        assert turns[0]["user"].content == "Q1"
        assert turns[0]["assistant"].content == "A1"

    def test_multiple_turns(self):
        msgs = [
            _make_msg("user", "Q1", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg("assistant", "A1", created_at=datetime(2026, 1, 1, 0, 1)),
            _make_msg("user", "Q2", created_at=datetime(2026, 1, 1, 0, 2)),
            _make_msg("assistant", "A2", created_at=datetime(2026, 1, 1, 0, 3)),
        ]
        turns = _pair_messages(msgs)
        assert len(turns) == 2

    def test_orphan_assistant(self):
        msgs = [_make_msg("assistant", "orphan")]
        turns = _pair_messages(msgs)
        assert len(turns) == 1
        assert turns[0]["user"] is None
        assert turns[0]["assistant"].content == "orphan"


# ===================================================================
# Unit tests: SSE event parsing
# ===================================================================


class TestExtractSseEvents:
    def test_no_metadata(self):
        msg = _make_msg("assistant", "hi", metadata_=None)
        assert _extract_sse_events(msg) == []

    def test_no_sse_key(self):
        msg = _make_msg("assistant", "hi", metadata_={"answer": "hi"})
        assert _extract_sse_events(msg) == []

    def test_with_events(self):
        events = [{"event": "done", "data": {}}]
        msg = _make_msg("assistant", "hi", metadata_={"sse_events": events})
        assert _extract_sse_events(msg) == events


class TestExtractReactSteps:
    def test_empty(self):
        assert _extract_react_steps([]) == []

    def test_filters_tool_start(self):
        events = [
            {"event": "step", "data": {"type": "tool_start", "tool_name": "web_search"}},
            {"event": "step", "data": {"type": "tool_call", "tool_name": "web_search", "observation": "ok"}},
        ]
        steps = _extract_react_steps(events)
        assert len(steps) == 1
        assert steps[0]["type"] == "tool_call"


class TestExtractDoneEvent:
    def test_missing(self):
        assert _extract_done_event([]) is None

    def test_found(self):
        events = [
            {"event": "step", "data": {}},
            {"event": "done", "data": {"answer": "result", "iterations": 2}},
        ]
        done = _extract_done_event(events)
        assert done is not None
        assert done["answer"] == "result"


class TestExtractDagPlan:
    def test_empty(self):
        assert _extract_dag_plan([]) == []

    def test_round_match(self):
        events = [
            {
                "event": "phase",
                "data": {
                    "name": "planning",
                    "status": "done",
                    "round": 1,
                    "steps": [{"id": "S1", "task": "Search", "deps": []}],
                },
            }
        ]
        plan = _extract_dag_plan(events, 1)
        assert len(plan) == 1
        assert plan[0]["id"] == "S1"

    def test_round_mismatch(self):
        events = [
            {
                "event": "phase",
                "data": {
                    "name": "planning",
                    "status": "done",
                    "round": 2,
                    "steps": [{"id": "S1", "task": "Search", "deps": []}],
                },
            }
        ]
        assert _extract_dag_plan(events, 1) == []


class TestExtractDagAnalysis:
    def test_missing(self):
        assert _extract_dag_analysis([]) is None

    def test_found(self):
        events = [
            {
                "event": "phase",
                "data": {
                    "name": "analyzing",
                    "status": "done",
                    "round": 1,
                    "achieved": True,
                    "confidence": 0.85,
                    "reasoning": "Done.",
                },
            }
        ]
        result = _extract_dag_analysis(events)
        assert result is not None
        assert result["achieved"] is True
        assert result["confidence"] == 0.85


class TestExtractDagRounds:
    def test_empty(self):
        assert _extract_dag_rounds([]) == [1]

    def test_multiple_rounds(self):
        events = [
            {"event": "phase", "data": {"round": 2}},
            {"event": "phase", "data": {"round": 1}},
            {"event": "phase", "data": {"round": 2}},
        ]
        assert _extract_dag_rounds(events) == [1, 2]


# ===================================================================
# Integration tests: Markdown rendering
# ===================================================================


class TestRenderMdReact:
    def test_summary_mode(self):
        conv = _make_conv()
        msgs = [
            _make_msg("user", "What is Python?", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "Python is a programming language.",
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        md = _render_md(conv, msgs, DetailLevel.SUMMARY)
        assert "# Test Conversation" in md
        assert "**User:**" in md
        assert "Python is a programming language." in md
        # Summary should NOT have execution details
        assert "Execution Details" not in md

    def test_full_mode_with_tool_calls(self):
        events = _react_events([
            {"tool": "web_search", "args": {"query": "python"}, "reasoning": "searching", "observation": "found", "elapsed": 1.5},
        ], answer="Python is great.")
        conv = _make_conv()
        msgs = [
            _make_msg("user", "Tell me about Python", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "Python is great.",
                metadata_={"sse_events": events, "answer": "Python is great."},
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        md = _render_md(conv, msgs, DetailLevel.FULL)
        assert "Execution Details" in md
        assert "web_search" in md
        assert "1.5s" in md
        assert "<details>" in md
        assert "Python is great." in md

    def test_empty_conversation(self):
        conv = _make_conv()
        md = _render_md(conv, [], DetailLevel.FULL)
        assert "# Test Conversation" in md
        # No turns
        assert "Turn 1" not in md


class TestRenderMdDag:
    def test_full_mode(self):
        plan = [
            {"id": "S1", "task": "Search for data", "deps": [], "tool_hint": "web_search"},
            {"id": "S2", "task": "Analyze results", "deps": ["S1"], "tool_hint": "python_exec"},
        ]
        events = _dag_events(plan, answer="DAG complete.")
        conv = _make_conv(mode="dag")
        msgs = [
            _make_msg("user", "Run a multi-step analysis", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "DAG complete.",
                metadata_={"sse_events": events, "answer": "DAG complete."},
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        md = _render_md(conv, msgs, DetailLevel.FULL)
        assert "Plan" in md
        assert "S1" in md
        assert "S2" in md
        assert "Search for data" in md
        assert "Analysis" in md
        assert "Goal Achieved" in md
        assert "90%" in md
        assert "DAG complete." in md


# ===================================================================
# Integration tests: TXT rendering
# ===================================================================


class TestRenderTxt:
    def test_summary(self):
        conv = _make_conv()
        msgs = [
            _make_msg("user", "Hello", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg("assistant", "Hi there!", created_at=datetime(2026, 1, 1, 0, 1)),
        ]
        txt = _render_txt(conv, msgs, DetailLevel.SUMMARY)
        assert "Test Conversation" in txt
        assert "[User]" in txt
        assert "[Assistant]" in txt
        assert "Hi there!" in txt
        # No markdown formatting
        assert "**" not in txt
        assert "#" not in txt

    def test_full_react(self):
        events = _react_events([
            {"tool": "calculator", "elapsed": 0.5},
        ])
        conv = _make_conv()
        msgs = [
            _make_msg("user", "2+2?", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "4",
                metadata_={"sse_events": events},
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        txt = _render_txt(conv, msgs, DetailLevel.FULL)
        assert "Execution Details" in txt
        assert "calculator" in txt


# ===================================================================
# Integration tests: DOCX rendering
# ===================================================================


class TestRenderDocx:
    def test_basic(self):
        conv = _make_conv()
        msgs = [
            _make_msg("user", "What is AI?", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg("assistant", "AI is artificial intelligence.", created_at=datetime(2026, 1, 1, 0, 1)),
        ]
        docx_bytes = _render_docx(conv, msgs, DetailLevel.SUMMARY)
        # DOCX files start with the PK zip header
        assert docx_bytes[:2] == b"PK"
        assert len(docx_bytes) > 100

    def test_full_with_react_steps(self):
        events = _react_events([
            {"tool": "web_search", "elapsed": 1.0},
        ])
        conv = _make_conv()
        msgs = [
            _make_msg("user", "Search something", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "Found it.",
                metadata_={"sse_events": events},
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        docx_bytes = _render_docx(conv, msgs, DetailLevel.FULL)
        assert docx_bytes[:2] == b"PK"

    def test_empty_conversation(self):
        conv = _make_conv()
        docx_bytes = _render_docx(conv, [], DetailLevel.FULL)
        assert docx_bytes[:2] == b"PK"

    def test_dag_full(self):
        plan = [{"id": "S1", "task": "Fetch data", "deps": [], "tool_hint": "web_fetch"}]
        events = _dag_events(plan, answer="Done.")
        conv = _make_conv(mode="dag")
        msgs = [
            _make_msg("user", "Fetch and analyze", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "Done.",
                metadata_={"sse_events": events},
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        docx_bytes = _render_docx(conv, msgs, DetailLevel.FULL)
        assert docx_bytes[:2] == b"PK"


# ===================================================================
# Edge cases
# ===================================================================


class TestEdgeCases:
    def test_assistant_with_answer_in_metadata_only(self):
        """When content is empty, fall back to metadata_.answer."""
        conv = _make_conv()
        msgs = [
            _make_msg("user", "Q", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "",
                metadata_={"answer": "Fallback answer", "sse_events": []},
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        md = _render_md(conv, msgs, DetailLevel.FULL)
        assert "Fallback answer" in md

    def test_no_model_name(self):
        """model_name should never appear in any export format."""
        conv = _make_conv(model_name="gpt-4o")  # even when set
        msgs = [
            _make_msg("user", "Hi", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg("assistant", "Hello", created_at=datetime(2026, 1, 1, 0, 1)),
        ]
        md = _render_md(conv, msgs, DetailLevel.SUMMARY)
        assert "Model" not in md
        assert "gpt-4o" not in md

        txt = _render_txt(conv, msgs, DetailLevel.SUMMARY)
        assert "Model" not in txt
        assert "gpt-4o" not in txt

        # DOCX -- parse the document to check
        from docx import Document

        docx_bytes = _render_docx(conv, msgs, DetailLevel.SUMMARY)
        doc = Document(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Model" not in full_text
        assert "gpt-4o" not in full_text

    def test_user_message_without_assistant(self):
        conv = _make_conv()
        msgs = [_make_msg("user", "Unanswered question")]
        md = _render_md(conv, msgs, DetailLevel.FULL)
        assert "Unanswered question" in md
        assert "Turn 1" in md

    def test_no_total_tokens_in_summary(self):
        conv = _make_conv(total_tokens=0)
        md = _render_md(conv, [], DetailLevel.SUMMARY)
        assert "Total Tokens" not in md


# ===================================================================
# Integration tests: DOCX markdown rendering
# ===================================================================


class TestDocxMarkdownRendering:
    """Verify that assistant's markdown answer renders as proper DOCX elements."""

    def test_headings_rendered(self):
        """Headings in markdown should become DOCX headings, not raw '## text'."""
        conv = _make_conv()
        msgs = [
            _make_msg("user", "Explain", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "# Title\n\n## Section\n\nSome text.",
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        docx_bytes = _render_docx(conv, msgs, DetailLevel.SUMMARY)
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        # Find heading paragraphs
        heading_texts = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        # The assistant's markdown headings should be real headings
        assert any("Title" in h for h in heading_texts)
        assert any("Section" in h for h in heading_texts)
        # Raw markdown markers should NOT appear
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "## Section" not in all_text

    def test_bold_rendered(self):
        """Bold markdown should become bold runs, not raw **text**."""
        conv = _make_conv()
        msgs = [
            _make_msg("user", "Q", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "This is **important** text.",
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        docx_bytes = _render_docx(conv, msgs, DetailLevel.SUMMARY)
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "**important**" not in all_text  # raw markers gone
        assert "important" in all_text  # content preserved
        # Check that at least one run is bold
        has_bold = any(
            run.bold for p in doc.paragraphs for run in p.runs if run.bold
        )
        assert has_bold

    def test_code_block_rendered(self):
        """Code blocks should use monospace font."""
        conv = _make_conv()
        answer = "Here is code:\n\n```python\nprint('hello')\n```\n\nDone."
        msgs = [
            _make_msg("user", "Show code", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg("assistant", answer, created_at=datetime(2026, 1, 1, 0, 1)),
        ]
        docx_bytes = _render_docx(conv, msgs, DetailLevel.SUMMARY)
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        # Code block should have Courier New font
        has_courier = any(
            run.font.name == "Courier New"
            for p in doc.paragraphs
            for run in p.runs
            if run.font.name == "Courier New"
        )
        assert has_courier

    def test_list_rendered(self):
        """Lists should use List Bullet/Number styles."""
        conv = _make_conv()
        answer = "Items:\n\n- First\n- Second\n- Third"
        msgs = [
            _make_msg("user", "List", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg("assistant", answer, created_at=datetime(2026, 1, 1, 0, 1)),
        ]
        docx_bytes = _render_docx(conv, msgs, DetailLevel.SUMMARY)
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        bullet_styles = [
            p.style.name for p in doc.paragraphs if "List" in p.style.name
        ]
        assert len(bullet_styles) >= 3  # at least 3 list items


# ===================================================================
# Integration tests: PDF rendering
# ===================================================================


class TestRenderPdf:
    def test_basic_pdf(self):
        """Basic PDF generation should produce valid PDF bytes."""
        conv = _make_conv()
        msgs = [
            _make_msg("user", "What is AI?", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "AI is artificial intelligence.",
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        pdf_bytes = _render_pdf(conv, msgs, DetailLevel.SUMMARY)
        assert pdf_bytes[:5] == b"%PDF-"
        assert len(pdf_bytes) > 100

    def test_summary_mode(self):
        """Summary PDF should not include execution details."""
        events = _react_events([{"tool": "web_search", "elapsed": 1.0}])
        conv = _make_conv()
        msgs = [
            _make_msg("user", "Search", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "Found it.",
                metadata_={"sse_events": events},
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        pdf_bytes = _render_pdf(conv, msgs, DetailLevel.SUMMARY)
        assert pdf_bytes[:5] == b"%PDF-"

    def test_full_react_mode(self):
        """Full mode PDF with ReAct steps should produce valid PDF."""
        events = _react_events(
            [
                {
                    "tool": "web_search",
                    "args": {"query": "test"},
                    "reasoning": "searching",
                    "observation": "found",
                    "elapsed": 1.5,
                },
            ],
            answer="Result.",
        )
        conv = _make_conv()
        msgs = [
            _make_msg(
                "user", "Search for info", created_at=datetime(2026, 1, 1, 0, 0)
            ),
            _make_msg(
                "assistant",
                "Result.",
                metadata_={"sse_events": events, "answer": "Result."},
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        pdf_bytes = _render_pdf(conv, msgs, DetailLevel.FULL)
        assert pdf_bytes[:5] == b"%PDF-"
        assert len(pdf_bytes) > 500

    def test_full_dag_mode(self):
        """Full mode PDF with DAG steps."""
        plan = [
            {"id": "S1", "task": "Search", "deps": [], "tool_hint": "web_search"},
            {
                "id": "S2",
                "task": "Analyze",
                "deps": ["S1"],
                "tool_hint": "python_exec",
            },
        ]
        events = _dag_events(plan, answer="DAG done.")
        conv = _make_conv(mode="dag")
        msgs = [
            _make_msg("user", "Multi-step", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg(
                "assistant",
                "DAG done.",
                metadata_={"sse_events": events, "answer": "DAG done."},
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        pdf_bytes = _render_pdf(conv, msgs, DetailLevel.FULL)
        assert pdf_bytes[:5] == b"%PDF-"

    def test_cjk_content(self):
        """PDF with Chinese content should not crash."""
        conv = _make_conv(title="测试对话")
        msgs = [
            _make_msg(
                "user", "你好，请解释人工智能", created_at=datetime(2026, 1, 1, 0, 0)
            ),
            _make_msg(
                "assistant",
                "人工智能是计算机科学的一个分支。",
                created_at=datetime(2026, 1, 1, 0, 1),
            ),
        ]
        pdf_bytes = _render_pdf(conv, msgs, DetailLevel.SUMMARY)
        assert pdf_bytes[:5] == b"%PDF-"

    def test_empty_conversation(self):
        """PDF with no messages should still produce a valid PDF."""
        conv = _make_conv()
        pdf_bytes = _render_pdf(conv, [], DetailLevel.FULL)
        assert pdf_bytes[:5] == b"%PDF-"

    def test_markdown_in_answer(self):
        """PDF should handle markdown content in the assistant answer."""
        conv = _make_conv()
        answer = (
            "# Hello\n\nThis is **bold** and *italic*.\n\n"
            "```python\nprint('hi')\n```\n\n- item 1\n- item 2"
        )
        msgs = [
            _make_msg(
                "user", "Show me stuff", created_at=datetime(2026, 1, 1, 0, 0)
            ),
            _make_msg("assistant", answer, created_at=datetime(2026, 1, 1, 0, 1)),
        ]
        pdf_bytes = _render_pdf(conv, msgs, DetailLevel.SUMMARY)
        assert pdf_bytes[:5] == b"%PDF-"
        assert len(pdf_bytes) > 200


# ===================================================================
# Unit tests: _strip_emoji
# ===================================================================


class TestStripEmoji:
    def test_emoji_removed(self):
        assert _strip_emoji("Hello \U0001F30D World \U0001F680") == "Hello  World "

    def test_bmp_emoji_removed(self):
        """Emoji from BMP blocks (Misc Technical, Geometric Shapes, etc.)."""
        assert _strip_emoji("⭐ star") == " star"         # U+2B50
        assert _strip_emoji("⏰ alarm") == " alarm"       # U+23F0
        assert _strip_emoji("▶ play") == " play"          # U+25B6
        assert _strip_emoji("ℹ info") == " info"          # U+2139

    def test_enclosed_supplement_removed(self):
        """Emoji from Enclosed Ideographic Supplement (U+1F200 block)."""
        assert _strip_emoji("🈁 koko") == " koko"         # U+1F201

    def test_normal_text_unchanged(self):
        assert _strip_emoji("Hello World") == "Hello World"

    def test_cjk_not_stripped(self):
        text = "你好世界 Hello"
        assert _strip_emoji(text) == text

    def test_empty_string(self):
        assert _strip_emoji("") == ""


# ===================================================================
# Integration tests: DOCX heading color
# ===================================================================


class TestDocxHeadingColor:
    def test_headings_use_amber_color(self):
        """DOCX headings should use amber color (0x946B2D), not default blue."""
        from docx import Document
        from docx.shared import RGBColor

        conv = _make_conv()
        msgs = [
            _make_msg("user", "Hi", created_at=datetime(2026, 1, 1, 0, 0)),
            _make_msg("assistant", "Hello", created_at=datetime(2026, 1, 1, 0, 1)),
        ]
        docx_bytes = _render_docx(conv, msgs, DetailLevel.SUMMARY)
        doc = Document(io.BytesIO(docx_bytes))

        expected = RGBColor(0x94, 0x6B, 0x2D)
        style = doc.styles["Heading 1"]
        assert style.font.color.rgb == expected
